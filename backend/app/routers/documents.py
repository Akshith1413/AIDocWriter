import io
import json
from html import escape

import bleach
import markdown
from docx import Document as WordDocument
from fastapi import APIRouter, Depends, HTTPException, Query, Response
from sqlalchemy import func, select
from sqlalchemy.orm import Session

from ..database import get_db
from ..models import Document, User
from ..schemas import DashboardSummary, DocumentUpdate, DocumentView, GenerateRequest
from ..security import get_current_user
from ..serializers import document_view
from ..templates import get_template
from ..workflow import DocumentOrchestrator

router = APIRouter(prefix="/documents", tags=["documents"])


def owned_document(document_id: str, user: User, db: Session) -> Document:
    document = db.scalar(
        select(Document).where(Document.id == document_id, Document.user_id == user.id)
    )
    if not document:
        raise HTTPException(status_code=404, detail="Document not found.")
    return document


@router.post("/generate", response_model=DocumentView, status_code=201)
async def generate_document(
    payload: GenerateRequest,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
) -> DocumentView:
    from ..llm import ProviderError
    try:
        result = await DocumentOrchestrator(payload).generate()
    except ProviderError as exc:
        raise HTTPException(status_code=502, detail=str(exc))
    document = Document(
        user_id=user.id,
        title=result.title,
        template=payload.template,
        source_notes=payload.input_text,
        content_md=result.content_md,
        review_json=result.review.model_dump_json(),
        provider=result.provider,
        model=result.model,
        status=result.status,
        iteration_count=result.iteration_count,
    )
    db.add(document)
    db.commit()
    db.refresh(document)
    return document_view(document)


@router.get("", response_model=list[DocumentView])
def list_documents(
    user: User = Depends(get_current_user), db: Session = Depends(get_db)
) -> list[DocumentView]:
    documents = db.scalars(
        select(Document).where(Document.user_id == user.id).order_by(Document.updated_at.desc())
    ).all()
    return [document_view(document) for document in documents]


@router.get("/dashboard", response_model=DashboardSummary)
def dashboard(
    user: User = Depends(get_current_user), db: Session = Depends(get_db)
) -> DashboardSummary:
    base = Document.user_id == user.id
    total = db.scalar(select(func.count(Document.id)).where(base)) or 0
    approved = db.scalar(
        select(func.count(Document.id)).where(base, Document.status == "approved")
    ) or 0
    average = db.scalar(select(func.avg(Document.iteration_count)).where(base)) or 0
    recent = db.scalars(
        select(Document).where(base).order_by(Document.updated_at.desc()).limit(5)
    ).all()
    return DashboardSummary(
        total_documents=total,
        approved_documents=approved,
        needs_attention=total - approved,
        average_iterations=round(float(average), 1),
        recent_documents=[document_view(document) for document in recent],
    )


@router.get("/{document_id}", response_model=DocumentView)
def get_document(
    document_id: str, user: User = Depends(get_current_user), db: Session = Depends(get_db)
) -> DocumentView:
    return document_view(owned_document(document_id, user, db))


@router.patch("/{document_id}", response_model=DocumentView)
def update_document(
    document_id: str,
    payload: DocumentUpdate,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
) -> DocumentView:
    document = owned_document(document_id, user, db)
    if payload.title is not None:
        document.title = payload.title.strip()
    if payload.content_md is not None:
        document.content_md = payload.content_md
        document.status = "revision_required"
    db.commit()
    db.refresh(document)
    return document_view(document)


@router.post("/{document_id}/review", response_model=DocumentView)
async def review_document(
    document_id: str,
    auto_refine: bool = Query(default=True),
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
) -> DocumentView:
    document = owned_document(document_id, user, db)
    from ..llm import ProviderError
    request = GenerateRequest(
        title=document.title,
        input_text=document.source_notes,
        template=document.template,
        provider=document.provider,
        model=document.model,
    )
    orchestrator = DocumentOrchestrator(request)
    try:
        if auto_refine:
            result = await orchestrator.review_existing(document.title, payload.draft)
        else:
            review = await orchestrator.client.review(
                get_template(document.template),
                document.source_notes,
                payload.draft,
            )
            document.content_md = payload.draft
            document.review_json = review.model_dump_json()
            document.status = review.status
            db.commit()
            db.refresh(document)
            return document_view(document)
    except ProviderError as exc:
        raise HTTPException(status_code=502, detail=str(exc))
    
    document.content_md = result.content_md
    document.review_json = result.review.model_dump_json()
    document.status = result.status
    document.iteration_count += result.iteration_count
    db.commit()
    db.refresh(document)
    return document_view(document)


@router.delete("/{document_id}", status_code=204)
def delete_document(
    document_id: str, user: User = Depends(get_current_user), db: Session = Depends(get_db)
) -> Response:
    document = owned_document(document_id, user, db)
    db.delete(document)
    db.commit()
    return Response(status_code=204)


def html_export(document: Document) -> bytes:
    body = markdown.markdown(document.content_md, extensions=["tables", "fenced_code"])
    clean_body = bleach.clean(
        body,
        tags=[
            "h1", "h2", "h3", "h4", "p", "strong", "em", "blockquote", "ul", "ol", "li",
            "table", "thead", "tbody", "tr", "th", "td", "code", "pre", "hr", "br",
        ],
        attributes={},
    )
    html = f"""<!doctype html>
<html><head><meta charset="utf-8"><title>{escape(document.title)}</title>
<style>body{{max-width:900px;margin:48px auto;padding:0 32px;font:16px/1.6 Arial;color:#182331}}
h1,h2{{color:#0e2948}}table{{border-collapse:collapse;width:100%}}th,td{{border:1px solid #cad5e2;padding:9px}}
blockquote{{border-left:4px solid #5b6eff;padding:8px 16px;background:#f4f6ff}}</style></head>
<body>{clean_body}</body></html>"""
    return html.encode("utf-8")


def word_export(document: Document) -> bytes:
    output = WordDocument()
    for line in document.content_md.splitlines():
        if line.startswith("# "):
            output.add_heading(line[2:], level=0)
        elif line.startswith("## "):
            output.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            output.add_heading(line[4:], level=2)
        elif line.startswith("> "):
            output.add_paragraph(line[2:], style="Quote")
        elif line.startswith("- "):
            output.add_paragraph(line[2:], style="List Bullet")
        elif line.strip() and not line.startswith("|"):
            output.add_paragraph(line)
    buffer = io.BytesIO()
    output.save(buffer)
    return buffer.getvalue()


@router.get("/{document_id}/export/{file_type}")
def export_document(
    document_id: str,
    file_type: str,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
) -> Response:
    document = owned_document(document_id, user, db)
    filename = "".join(char if char.isalnum() or char in "-_" else "-" for char in document.title)[:80]
    if file_type == "md":
        return Response(
            document.content_md,
            media_type="text/markdown",
            headers={"Content-Disposition": f'attachment; filename="{filename}.md"'},
        )
    if file_type == "json":
        payload = json.dumps(document_view(document).model_dump(mode="json"), indent=2)
        return Response(
            payload,
            media_type="application/json",
            headers={"Content-Disposition": f'attachment; filename="{filename}.json"'},
        )
    if file_type == "html":
        return Response(
            html_export(document),
            media_type="text/html",
            headers={"Content-Disposition": f'attachment; filename="{filename}.html"'},
        )
    if file_type == "docx":
        return Response(
            word_export(document),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{filename}.docx"'},
        )
    raise HTTPException(status_code=400, detail="Supported exports are md, json, html, and docx.")
